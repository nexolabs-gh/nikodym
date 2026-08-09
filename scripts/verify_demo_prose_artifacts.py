"""Verifica prosa y resumen PSI en todos los formatos finales de una familia demo.

Este gate abre los bytes que consume la persona: HTML, PDF, Word, QMD dentro del ZIP y JSON. No
compara hashes de formatos con metadatos variables; normaliza texto y comprueba afirmaciones
semánticas. Los capturadores lo invocan por familia después de escribir de forma atómica.
"""

from __future__ import annotations

import argparse
import html
import json
import re
import unicodedata
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from docx import Document
from pypdf import PdfReader

FamilyName = Literal["f1", "f3", "ifrs9"]

_ROOT = Path(__file__).resolve().parent.parent
_FIXTURES = _ROOT / "web" / "src" / "fixtures" / "demo"
_COMMON_REQUIRED = ("se reporta para revisión toda variable con más de",)
_COMMON_FORBIDDEN = ("se rechaza toda variable",)
_SCORECARD_REQUIRED = (
    "El peor PSI entre score y PD en",
    "IV igual o superior a",
    "Los tramos efectivos por partición fueron",
    "ordenados por la PD calibrada",
)
_SCORECARD_FORBIDDEN = (
    "El PSI del score en Desarrollo vs.",
    "IV superior a",
    "reparte la población en 10 tramos",
    "ordenados por score",
    "Estable ≤",
    "Revisar ≤",
)


@dataclass(frozen=True)
class _FamilyPaths:
    html: Path
    pdf: Path
    docx: Path
    quarto: Path
    results: Path


def verify_demo_family(family: FamilyName, *, fixtures: Path = _FIXTURES) -> None:
    """Abre y verifica los cinco entregables canónicos de una familia demo."""
    paths = _family_paths(family, fixtures)
    texts = {
        "HTML": _html_text(paths.html),
        "PDF": _pdf_text(paths.pdf),
        "Word": _docx_text(paths.docx),
        "QMD": _qmd_text(paths.quarto),
    }
    required = _COMMON_REQUIRED + (_SCORECARD_REQUIRED if family in {"f1", "f3"} else ())
    forbidden = _COMMON_FORBIDDEN + (_SCORECARD_FORBIDDEN if family in {"f1", "f3"} else ())
    for format_name, text in texts.items():
        for phrase in required:
            assert _normalize(phrase) in text, (
                f"{family}/{format_name}: falta la afirmación factual {phrase!r}."
            )
        for phrase in forbidden:
            assert _normalize(phrase) not in text, (
                f"{family}/{format_name}: reapareció la prosa falsa {phrase!r}."
            )

    results = json.loads(paths.results.read_text(encoding="utf-8"))
    if family in {"f1", "f3"}:
        _verify_psi_summary(results, family=family)


def verify_all_demo_families(*, fixtures: Path = _FIXTURES) -> None:
    """Verifica F1, F3 e IFRS 9 contra el mismo oráculo de artefactos."""
    for family in ("f1", "f3", "ifrs9"):
        verify_demo_family(family, fixtures=fixtures)


def _verify_psi_summary(results: object, *, family: FamilyName) -> None:
    assert isinstance(results, dict), f"{family}/JSON: la raíz no es un objeto."
    stability = results.get("stability")
    assert isinstance(stability, dict), f"{family}/JSON: falta la card de stability."
    rows = stability.get("stability_metrics")
    assert isinstance(rows, list), f"{family}/JSON: falta stability_metrics."
    maxima = stability.get("max_psi_by_comparison")
    metrics = stability.get("psi_metric_by_comparison")
    bands = stability.get("bands_by_comparison")
    comparisons = stability.get("comparisons")
    assert isinstance(maxima, dict) and isinstance(metrics, dict) and isinstance(bands, dict)
    assert isinstance(comparisons, list)
    assert set(maxima) == set(metrics) == set(bands) == set(comparisons)

    for comparison in comparisons:
        candidates = [
            row
            for row in rows
            if isinstance(row, dict)
            and row.get("comparison") == comparison
            and row.get("metric") in {"score_psi", "pd_psi"}
            and isinstance(row.get("value"), (int, float))
            and not isinstance(row.get("value"), bool)
        ]
        if not candidates:
            assert maxima[comparison] is None
            assert metrics[comparison] is None
            assert bands[comparison] == "not_evaluable"
            continue
        winner = max(
            candidates,
            key=lambda row: (float(row["value"]), row["metric"] == "score_psi"),
        )
        assert maxima[comparison] == winner["value"], (
            f"{family}/JSON/{comparison}: el máximo no pertenece a la fila ganadora."
        )
        assert metrics[comparison] == winner["metric"], (
            f"{family}/JSON/{comparison}: la identidad no pertenece a la fila ganadora."
        )
        assert bands[comparison] == winner["band"], (
            f"{family}/JSON/{comparison}: la banda no pertenece a la fila ganadora."
        )


def _family_paths(family: FamilyName, fixtures: Path) -> _FamilyPaths:
    suffix = {"f1": "-f1", "f3": "", "ifrs9": "-ifrs9"}[family]
    quarto = fixtures / ("report-quarto.zip" if family == "f3" else f"report-quarto{suffix}.zip")
    return _FamilyPaths(
        html=fixtures / f"report{suffix}.html",
        pdf=fixtures / f"report{suffix}.pdf",
        docx=fixtures / f"report{suffix}.docx",
        quarto=quarto,
        results=fixtures / f"results{suffix}.json",
    )


def _html_text(path: Path) -> str:
    raw = path.read_text(encoding="utf-8")
    return _normalize(re.sub(r"<[^>]+>", " ", html.unescape(raw)))


def _pdf_text(path: Path) -> str:
    return _normalize("\n".join(page.extract_text() or "" for page in PdfReader(path).pages))


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        parts.extend(cell.text for row in table.rows for cell in row.cells)
    return _normalize("\n".join(parts))


def _qmd_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return _normalize(archive.read("report.qmd").decode("utf-8"))


def _normalize(value: str) -> str:
    normalized = unicodedata.normalize("NFKC", value).replace("\u00a0", " ")
    return " ".join(normalized.split())


def main() -> None:
    """Ejecuta el verificador para una familia o para las tres."""
    parser = argparse.ArgumentParser()
    parser.add_argument("family", choices=("f1", "f3", "ifrs9", "all"), nargs="?", default="all")
    args = parser.parse_args()
    if args.family == "all":
        verify_all_demo_families()
    else:
        verify_demo_family(args.family)


if __name__ == "__main__":
    main()
