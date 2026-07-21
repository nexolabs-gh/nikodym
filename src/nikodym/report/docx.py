"""Export ``.docx`` (Word) del informe vía ``python-docx``, con import perezoso (mejora 1.1).

Word es lo que Validación **realmente edita** en un banco. El ``.qmd`` sirve al analista que
compila; el ``.docx`` sirve al comité que comenta, marca y firma. Por eso usa estilos nativos
—``Heading 1/2/3`` reales, de modo que el índice automático de Word funcione— y **tablas nativas**,
no imágenes de tablas: una tabla que no se puede editar ni copiar no es un documento de trabajo.

Sigue el patrón de :mod:`nikodym.report.pdf`: la dependencia (``python-docx``, extra ``docx``) se
importa **dentro** de la función, nunca en import-time, y su ausencia se traduce a un
:class:`~nikodym.report.exceptions.ReportDependencyError` accionable. A diferencia del PDF,
``python-docx`` es MIT puro (WeasyPrint arrastra Pyphen, que roza GPL), así que el extra ``docx``
**sí** entra en el meta-extra ``all`` y en el cierre redistribuible.

El ``.docx`` es un ZIP y **no** es byte-determinista (guarda marcas de tiempo de compresión); su
verificación es estructural —abre en Word sin reparar, con sus encabezados y tablas—, no por digest,
igual que el PDF. El *contenido* sí es determinista: sale de la misma
:func:`~nikodym.report.renderer.build_document_view` que el HTML y el ``.qmd``.

**Experimental (fuera de la garantía SemVer 1.x).**
"""

from __future__ import annotations

import io
import warnings
from collections.abc import Mapping, Sequence
from pathlib import Path
from typing import Any, Final

from nikodym.report.config import ReportConfig
from nikodym.report.exceptions import ReportDependencyError, ReportExportError
from nikodym.report.renderer import build_document_view
from nikodym.report.results import AiNarrationBlock, ReportInputBundle

__all__ = ["DocxReportRenderer", "render_docx"]

_DOCX_EXTENSION: Final = ".docx"
_TABLE_STYLE: Final = "Table Grid"
_MONO_FONT: Final = "Consolas"
# Ámbar suave: el bloque POR COMPLETAR tiene que saltar a la vista al abrir el archivo.
_PLACEHOLDER_FILL: Final = "FFF3CD"
_MISSING_FILL: Final = "F8D7DA"
_FIGURE_WIDTH_INCHES: Final = 6.0
# Techo de filas por tabla en Word. El detalle por observación ya no vive en el documento (va como
# adjunto), así que ninguna tabla real se acerca a este límite; es un cinturón, no una política.
_MAX_DOCX_TABLE_ROWS: Final = 500


def render_docx(document: Mapping[str, Any], *, config: ReportConfig) -> bytes:
    """Renderiza la vista del documento a bytes ``.docx`` (import perezoso de ``python-docx``).

    Parameters
    ----------
    document:
        Proyección canónica devuelta por
        :func:`~nikodym.report.renderer.build_document_view` con ``chart_format='png'``: Word no
        admite figuras SVG.
    config:
        Sección ``report`` ya validada; aporta los metadatos de portada.

    Returns
    -------
    bytes
        Contenido del ``.docx`` (un ZIP OOXML). No es byte-determinista entre corridas.

    Raises
    ------
    ReportDependencyError
        Si ``python-docx`` no está instalado.
    """
    try:
        import docx
        from docx.shared import Inches, Pt
    except ModuleNotFoundError as exc:
        raise ReportDependencyError(
            "No se pudo generar el .docx: falta python-docx. Instale `nikodym[docx]` "
            "(o `nikodym[all]`, que ya lo incluye) y reintente."
        ) from exc

    word = docx.Document()
    _set_core_properties(word, document, config)

    word.add_heading(str(document["document_title"]), level=0)
    _cover(word, document, config)
    _toc_field(word)
    _executive(word, document["executive"])
    for section in document["sections"]:
        _section(word, section, config, inches=Inches, points=Pt)

    buffer = io.BytesIO()
    word.save(buffer)
    return buffer.getvalue()


class DocxReportRenderer:
    """Render opcional a ``.docx`` sobre la misma proyección de documento que el HTML."""

    def __init__(self, config: ReportConfig | None = None) -> None:
        """Construye el renderer desde ``ReportConfig`` (o el default, para uso standalone)."""
        self.config = config if config is not None else ReportConfig()

    @classmethod
    def from_config(cls, cfg: ReportConfig) -> DocxReportRenderer:
        """Construye ``DocxReportRenderer`` desde ``NikodymConfig.report``."""
        return cls(cfg)

    def render(
        self,
        bundle: ReportInputBundle,
        *,
        ai_blocks: tuple[AiNarrationBlock, ...] = (),
    ) -> bytes:
        """Renderiza el ``.docx`` desde el bundle; los gráficos van en PNG (Word no admite SVG)."""
        document = build_document_view(
            bundle, config=self.config, ai_blocks=ai_blocks, chart_format="png"
        )
        return render_docx(document, config=self.config)

    def write_docx_from_bundle(
        self,
        bundle: ReportInputBundle,
        *,
        output_dir: str,
        ai_blocks: tuple[AiNarrationBlock, ...] = (),
    ) -> Path | None:
        """Escribe ``{basename}.docx`` y devuelve su ``Path``, o ``None`` al degradar con gracia.

        Espejo de :meth:`~nikodym.report.renderer.PdfReportRenderer.write_pdf_from_html`: sin
        ``python-docx``, re-lanza la dependencia (``docx.fail_if_unavailable=True``) o emite un
        ``RuntimeWarning`` y devuelve ``None`` (``False``), sin tumbar la corrida. El warning
        **propaga el diagnóstico** de la dependencia (qué extra instalar) en vez de reformularlo.
        """
        try:
            payload = self.render(bundle, ai_blocks=ai_blocks)
        except ReportDependencyError as exc:
            if self.config.docx.fail_if_unavailable:
                raise
            warnings.warn(
                f"{exc} — El reporte se generó sin el export .docx.",
                RuntimeWarning,
                stacklevel=2,
            )
            return None
        return self._write(payload, output_dir=output_dir)

    def _write(self, payload: bytes, *, output_dir: str) -> Path:
        """Escribe el ``.docx`` con escritura atómica (tmp + ``replace``), como el HTML y el PDF."""
        directory = Path(output_dir)
        directory.mkdir(parents=True, exist_ok=True)
        filename = f"{self.config.basename}{_DOCX_EXTENSION}"
        output_path = directory / filename
        temp_path = directory / f".{filename}.tmp"
        try:
            temp_path.write_bytes(payload)
            temp_path.replace(output_path)
        except OSError as exc:
            temp_path.unlink(missing_ok=True)
            raise ReportExportError(
                "No se pudo escribir el reporte Word: dominio='report', formato='docx', "
                f"clave='{filename}', output_dir='{output_dir}', "
                "acción='verifique permisos y espacio disponible'."
            ) from exc
        return output_path


# ─────────────────────────── bloques del documento Word ───────────────────────────


def _set_core_properties(word: Any, document: Mapping[str, Any], config: ReportConfig) -> None:
    """Rellena las propiedades del archivo Word desde el config y el lineage, no desde el reloj."""
    properties = word.core_properties
    document_title = str(document["document_title"])
    properties.title = document_title
    properties.subject = config.document.model_name or document_title
    properties.author = config.document.author or "Nikodym"
    properties.category = config.document.portfolio
    properties.comments = (
        f"Generado por Nikodym · config_hash={document['lineage']['config_hash']} · "
        f"data_hash={document['lineage']['data_hash']}"
    )


def _cover(word: Any, document: Mapping[str, Any], config: ReportConfig) -> None:
    """Portada: metadatos declarados, trazabilidad y los tres roles que firman."""
    del config
    subtitle = word.add_paragraph("USO INTERNO — CONFIDENCIAL")
    subtitle.runs[0].bold = True

    _key_value_table(
        word,
        [(field["label"], field["value"] if field["filled"] else "") for field in document["cover"]]
        + [("Fecha de emisión", document["emitted_date"])],
    )

    word.add_paragraph()
    caption = word.add_paragraph("Trazabilidad de la corrida (lineage)")
    caption.runs[0].bold = True
    lineage = document["lineage"]
    for key in ("config_hash", "data_hash", "git_sha", "root_seed"):
        _mono_paragraph(word, f"{key}={lineage[key]}")

    word.add_paragraph()
    _table(
        word,
        ("Rol", "Nombre", "Fecha"),
        [
            ("Desarrollo", "", ""),
            ("Validación independiente", "", ""),
            ("Aprobación", "", ""),
        ],
    )
    word.add_page_break()


def _toc_field(word: Any) -> None:
    """Inserta el campo TOC nativo de Word: el índice se arma solo desde los estilos de heading.

    No es un índice escrito a mano —que quedaría desincronizado en cuanto el validador añada un
    párrafo—: es el campo ``TOC`` de Word, que se actualiza con «Actualizar campos» y refleja
    siempre los encabezados reales del documento.
    """
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement

    word.add_heading("Índice", level=1)
    paragraph = word.add_paragraph()
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = r'TOC \o "1-3" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Índice: clic derecho sobre este texto → «Actualizar campos» en Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instruction, separate, placeholder, end):
        run._r.append(element)
    word.add_page_break()


def _executive(word: Any, executive: Mapping[str, Any]) -> None:
    """Resumen ejecutivo: el veredicto (que firma un humano) y las métricas clave."""
    word.add_heading("Resumen ejecutivo", level=1)
    _shaded_paragraph(
        word,
        "POR COMPLETAR — Veredicto de validación: Aprobado · Aprobado con observaciones · "
        "Rechazado. Campo estructural: lo marca y firma el validador independiente; no es un "
        "resultado calculado por el motor.",
        fill=_PLACEHOLDER_FILL,
    )
    metrics = executive["metrics"]
    if metrics:
        _table(
            word,
            ("Métrica", "Alcance", "Valor", "Estado"),
            [
                (metric["label"], metric["scope"], metric["value"], metric["band"])
                for metric in metrics
            ],
        )
    else:
        word.add_paragraph("Las métricas clave no están disponibles en esta corrida.")
    for note in executive["notes"]:
        word.add_paragraph(note).runs[0].italic = True


def _section(
    word: Any,
    section: Mapping[str, Any],
    config: ReportConfig,
    *,
    inches: Any,
    points: Any,
) -> None:
    """Escribe una sección: encabezado con estilo nativo, prosa, figuras, tablas y anexos."""
    del config
    if section["kind"] == "toc":
        return  # el índice de Word es el campo TOC nativo, no una lista escrita a mano.

    word.add_heading(_heading(section), level=1 if section["level"] == 1 else 2)

    if section["status"] == "missing":
        _shaded_paragraph(
            word,
            "Sección requerida ausente; el reporte parcial no inventa números ni oculta la "
            "ausencia.",
            fill=_MISSING_FILL,
        )
    if section["kind"] == "appendix" and section["level"] == 2:
        _mono_paragraph(word, f"Artefacto de origen: {section['source']}")
    for paragraph in section["body"]:
        word.add_paragraph(paragraph)
    if section["placeholder"]:
        placeholder = section["placeholder"]
        _shaded_paragraph(
            word, f"POR COMPLETAR — {placeholder['title']}", fill=_PLACEHOLDER_FILL, bold=True
        )
        for line in placeholder["guidance"]:
            _shaded_paragraph(word, line, fill=_PLACEHOLDER_FILL)
    for chart in section["charts"]:
        word.add_picture(io.BytesIO(chart["png"]), width=inches(_FIGURE_WIDTH_INCHES))
        _caption(word, chart["title"], points=points)
    for table in section["tables"]:
        _caption(word, table["title"], points=points, bold=True)
        _table(word, tuple(table["columns"]), [tuple(row) for row in table["rows"]])
        if table["truncated"]:
            _caption(
                word,
                f"… (mostrando {table['shown_rows']} de {table['total_rows']} filas)",
                points=points,
            )
    if section["data_exports"]:
        _data_exports(word, section["data_exports"])
    _payload_blocks(word, section)
    if section["narration"]:
        _narration(word, section["narration"])


def _heading(section: Mapping[str, Any]) -> str:
    """Texto del encabezado con la numeración del documento (idéntica al HTML y al ``.qmd``)."""
    number = str(section["number"])
    if number and section["kind"] == "appendix" and section["level"] == 1:
        number = f"Anexo {number} —"
    prefix = f"{number} " if number else ""
    return f"{prefix}{section['title']}"


def _data_exports(word: Any, exports: Mapping[str, Any]) -> None:
    """Declara dónde quedó el detalle por observación (los adjuntos de datos)."""
    word.add_heading("Detalle por observación", level=3)
    if exports["requested"]:
        word.add_paragraph(
            "Las tablas con una fila por operación no se reproducen en el informe: se entregan "
            "completas como archivos adjuntos a este documento."
        )
        _table(
            word,
            ("Tabla", "Filas", "Archivo adjunto"),
            [
                (
                    item["title"],
                    item["rows"],
                    item["filename"] + (f" — hoja {item['sheet']}" if item["sheet"] else ""),
                )
                for item in exports["attachments"]
            ],
        )
        return
    word.add_paragraph(
        "Esta corrida produjo tablas con una fila por operación que no se reproducen en el informe "
        "(son datos, no evidencia de validación) y que no se exportaron: añada 'csv' o 'xlsx' a "
        "report.formats para obtenerlas completas."
    )
    for item in exports["pending"]:
        word.add_paragraph(f"{item['title']} — {item['rows']} filas", style="List Bullet")


def _payload_blocks(word: Any, section: Mapping[str, Any]) -> None:
    """Payload y métricas del anexo: JSON en monoespaciada, legible y copiable."""
    for heading, items in (
        ("Parámetros y payload", section["payload_items"]),
        ("Métricas estructuradas", section["metric_items"]),
    ):
        if not items:
            continue
        word.add_heading(heading, level=3)
        for item in items:
            key = word.add_paragraph()
            key.add_run(str(item["key"])).bold = True
            _mono_paragraph(word, str(item["value"]))


def _narration(word: Any, narration: Mapping[str, Any]) -> None:
    """Narrativa IA opcional, siempre etiquetada como generada por IA."""
    if narration["label"]:
        label = word.add_paragraph()
        label.add_run(narration["label"]).bold = True
    if narration["text"]:
        word.add_paragraph(narration["text"])
    if narration["warning"]:
        word.add_paragraph(narration["warning"]).runs[0].italic = True


# ─────────────────────────── primitivas de Word ───────────────────────────


def _table(word: Any, columns: Sequence[str], rows: Sequence[Sequence[str]]) -> Any:
    """Tabla NATIVA de Word (no una imagen): editable, copiable y con encabezado en negrita."""
    visible = list(rows)[:_MAX_DOCX_TABLE_ROWS]
    table = word.add_table(rows=1, cols=len(columns))
    table.style = _TABLE_STYLE
    for index, column in enumerate(columns):
        cell = table.rows[0].cells[index]
        cell.text = str(column)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in visible:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    return table


def _key_value_table(word: Any, items: Sequence[tuple[str, str]]) -> None:
    """Tabla de dos columnas para los metadatos de portada; lo no declarado queda en blanco."""
    table = word.add_table(rows=0, cols=2)
    table.style = _TABLE_STYLE
    for label, value in items:
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value


def _mono_paragraph(word: Any, text: str) -> Any:
    """Párrafo en fuente monoespaciada: hashes, rutas y JSON se leen así, no en proporcional."""
    paragraph = word.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = _MONO_FONT
    return paragraph


def _shaded_paragraph(word: Any, text: str, *, fill: str, bold: bool = False) -> Any:
    """Párrafo sombreado: es lo que hace que un POR COMPLETAR salte a la vista al abrir Word."""
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement

    paragraph = word.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shading)
    return paragraph


def _caption(word: Any, text: str, *, points: Any, bold: bool = False) -> Any:
    """Pie/título de figura o tabla, en cuerpo menor."""
    paragraph = word.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = points(9)
    return paragraph
