"""Tests de ``report.docx``: el Word que Validación realmente edita.

Los tests que EJECUTAN ``python-docx`` van gateados con ``skipif`` (extra ``docx``), igual que
``test_ml_backends`` y ``test_report_pdf``: el job mínimo del CI instala sin extras y correrlos
incondicionalmente lo reventaría. Los que cubren la AUSENCIA de la dependencia (bloqueando su
import) y el import liviano del paquete corren SIEMPRE.

El ``.docx`` es un ZIP y no es byte-determinista (guarda marcas de tiempo), así que se verifica su
ESTRUCTURA: que reabra sin reparar, que use estilos ``Heading`` reales —de los que Word deriva su
índice—, que las tablas sean nativas (no imágenes) y que los bloques POR COMPLETAR estén sombreados.
"""

from __future__ import annotations

import builtins
import importlib.util
import io
import subprocess
import sys
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
import pytest

from nikodym.core.config import NikodymConfig
from nikodym.core.lineage import LineageBundle
from nikodym.core.study import Study
from nikodym.report.config import DocumentStructureConfig, DocxRenderConfig, ReportConfig
from nikodym.report.docx import DocxReportRenderer
from nikodym.report.exceptions import ReportDependencyError
from nikodym.report.results import PlaceholderBlock, ReportInputBundle, ReportSection
from nikodym.report.step import REPORT_REQUIRED_CARDS, ReportStep

_HAS_DOCX = importlib.util.find_spec("docx") is not None
_SKIP_DOCX = pytest.mark.skipif(not _HAS_DOCX, reason="requiere el extra docx (python-docx)")


def _lineage() -> LineageBundle:
    """Lineage fijo del bundle sintético."""
    return LineageBundle(
        git_sha="abc123",
        git_dirty=False,
        data_hash="data123456789abcdef",
        config_hash="cfg123456789abcdef",
        root_seed=42,
        uv_lock_hash="uv123",
        library_versions={"nikodym": "1.0.0"},
        determinism_caveats=[],
        created_at=datetime(2026, 6, 24, 9, 30, tzinfo=UTC),
        schema_version="1.0.0",
    )


def _bundle() -> ReportInputBundle:
    """Documento mínimo: un capítulo con placeholder, una subsección con tabla y un anexo."""
    sections = (
        ReportSection(
            id="introduction",
            title="Introducción",
            status="included",
            source_domain="report",
            source_key="introduction",
            kind="prose",
            number="1",
            body=("El informe documenta la validación del scorecard.",),
            placeholder=PlaceholderBlock(
                title="Introducción",
                guidance=("Declara el propósito y el alcance.",),
            ),
        ),
        ReportSection(
            id="results.model",
            title="Modelo PD",
            status="included",
            source_domain="model",
            source_key="model_card",
            kind="data",
            level=2,
            number="4.3",
            body=("El modelo final incluye 2 variables.",),
        ),
        ReportSection(
            id="appendix_tables",
            title="Tablas detalladas",
            status="included",
            source_domain="report",
            source_key="appendix_tables",
            kind="appendix",
            number="B",
        ),
    )
    return ReportInputBundle(
        lineage=_lineage(),
        cards={"model": {"selected_features": ("mora", "saldo")}},
        tables={
            "model.coefficients": pd.DataFrame({"feature": ["mora"], "beta": [1.25]}),
            "scorecard.score": pd.DataFrame(
                {"score": [640, 712]}, index=pd.Index(["op-000", "op-001"], name="loan_id")
            ),
        },
        figures={},
        sections=sections,
        missing_sections=(),
    )


def _bundle_validation() -> ReportInputBundle:
    """Bundle con validation para comprobar el formato Word compartido."""
    base = _bundle()
    sections = (
        *base.sections,
        ReportSection(
            id="validation",
            title="Validación formal",
            status="included",
            source_domain="report",
            source_key="validation",
            kind="prose",
            number="5",
            body=("El estado técnico agregado es pass.",),
            placeholder=PlaceholderBlock(
                title="Veredicto de validación formal",
                guidance=("Firmar el juicio humano.",),
            ),
        ),
        ReportSection(
            id="validation.calibration",
            title="Calibración",
            status="included",
            source_domain="validation",
            source_key="result",
            kind="data",
            level=2,
            number="5.1",
        ),
    )
    return base.model_copy(
        update={
            "cards": {
                **base.cards,
                "validation": {
                    "model_ref": "scorecard@oracle",
                    "families_run": ("calibration",),
                    "overall_status": "pass",
                    "n_tests": 1,
                    "n_failed": 0,
                },
            },
            "results": {"validation": {"oracle": "atomic-result"}},
            "tables": {
                **base.tables,
                "validation.calibration": pd.DataFrame(
                    {"test": ["hosmer_lemeshow"], "p_value": [0.48], "decision": ["pass"]}
                ),
            },
            "sections": sections,
        },
        deep=True,
    )


def _config(**kwargs: Any) -> ReportConfig:
    """Config con metadatos de portada declarados."""
    base: dict[str, Any] = {
        "document": DocumentStructureConfig(
            model_name="Scorecard consumo",
            entity="Banco Ejemplo S.A.",
            author="Riesgo de Crédito",
        )
    }
    base.update(kwargs)
    return ReportConfig(**base)


# ─────────────────────────── (a) .docx real (skipif python-docx) ───────────────────────────


@_SKIP_DOCX
def test_docx_abre_sin_reparar_y_usa_estilos_nativos_de_word(tmp_path: Path) -> None:
    """El ``.docx`` reabre limpio, con ``Heading`` reales y tablas NATIVAS (no imágenes).

    Es el criterio de aceptación real: Validación abre el archivo en Word, ve el índice automático
    (que Word deriva de los estilos ``Heading``) y **edita las tablas**. Una tabla renderizada como
    imagen sería un documento muerto.
    """
    import docx

    path = DocxReportRenderer.from_config(_config()).write_docx_from_bundle(
        _bundle(), output_dir=str(tmp_path)
    )

    assert path == tmp_path / "scorecard_report.docx"
    assert path.is_file()

    # (1) ZIP OOXML íntegro: si estuviera corrupto, Word pediría "reparar".
    with zipfile.ZipFile(path) as archivo:
        assert archivo.testzip() is None
        assert "[Content_Types].xml" in archivo.namelist()
        assert "word/document.xml" in archivo.namelist()

    # (2) Reabre con python-docx (round-trip completo del XML).
    word = docx.Document(str(path))
    estilos = {parrafo.style.name for parrafo in word.paragraphs}
    assert "Title" in estilos
    assert "Heading 1" in estilos  # capítulos
    assert "Heading 2" in estilos  # subsecciones → el índice de Word funciona

    # (3) Tablas nativas y editables, con su encabezado.
    assert len(word.tables) >= 2
    encabezados = [tabla.rows[0].cells[0].text for tabla in word.tables]
    assert "feature" in encabezados

    # (4) Metadatos del archivo (los que Word muestra en Propiedades).
    assert word.core_properties.title == "Informe de Validación de Scorecard"
    assert word.core_properties.author == "Riesgo de Crédito"


@_SKIP_DOCX
def test_docx_incluye_validacion_formal_tabla_y_veredicto_humano() -> None:
    """Word espeja el capítulo formal y mantiene su tabla editable y el juicio humano visible."""
    import docx

    payload = DocxReportRenderer.from_config(_config()).render(_bundle_validation())
    word = docx.Document(io.BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in word.paragraphs)
    headings = [
        paragraph.text
        for paragraph in word.paragraphs
        if paragraph.style.name in {"Heading 1", "Heading 2"}
    ]

    assert "5 Validación formal" in headings
    assert "5.1 Calibración" in headings
    assert "POR COMPLETAR — Veredicto de validación formal" in text
    assert any(table.rows[1].cells[0].text == "hosmer_lemeshow" for table in word.tables)


@_SKIP_DOCX
def test_docx_marca_los_por_completar_e_inserta_el_indice_de_word() -> None:
    """Los POR COMPLETAR van sombreados y el índice es el campo TOC nativo, no una lista a mano."""
    payload = DocxReportRenderer.from_config(_config()).render(_bundle())
    xml = zipfile.ZipFile(io.BytesIO(payload)).read("word/document.xml").decode("utf-8")

    assert 'w:fill="FFF3CD"' in xml  # sombreado del bloque POR COMPLETAR: salta a la vista
    assert "POR COMPLETAR — Introducción" in xml
    assert "Declara el propósito y el alcance." in xml
    assert 'TOC \\o "1-3"' in xml  # campo TOC de Word: se actualiza solo desde los headings
    assert 'w:tblStyle w:val="TableGrid"' in xml  # tablas con bordes reales


@_SKIP_DOCX
def test_docx_no_arrastra_el_dataset_pero_dice_donde_esta() -> None:
    """Las tablas por observación no entran al Word: se referencian como adjunto."""
    renderer = DocxReportRenderer.from_config(_config(formats=("docx", "csv")))
    xml = zipfile.ZipFile(io.BytesIO(renderer.render(_bundle()))).read("word/document.xml").decode()

    assert "op-000" not in xml  # ninguna fila del frame por observación
    assert "Detalle por observación" in xml
    assert "scorecard_report__scorecard_score.csv" in xml


@_SKIP_DOCX
def test_step_con_formats_docx_escribe_el_word_real(tmp_path: Path) -> None:
    """``ReportStep`` con ``"docx"`` en ``formats`` escribe el archivo y lo refleja en el result.

    Cablea el objetivo de punta a punta: ``formats`` pide el Word, el step lo genera del MISMO
    documento que el HTML y ``ReportResult.docx_path`` apunta al archivo en disco.
    """
    import docx

    cfg = ReportConfig(output_dir=str(tmp_path), formats=("html", "docx"))
    study = _study_with_cards(cfg)

    result = ReportStep.from_config(cfg).execute(study, np.random.default_rng(20_240_629))

    assert result.docx_path == str(tmp_path / "scorecard_report.docx")
    assert Path(result.docx_path).is_file()
    assert docx.Document(result.docx_path).paragraphs  # reabre sin reparar
    assert result.manifest.output_format == "html"  # el .docx NO entra al manifest, como el PDF


# ─────────────────── (b) SIN python-docx: degradación (corre SIEMPRE) ───────────────────


def _bloquear_python_docx(monkeypatch: pytest.MonkeyPatch) -> None:
    """Fuerza ``ModuleNotFoundError`` al importar ``docx`` (ausencia determinista del extra)."""
    real_import = builtins.__import__

    def fake_import(
        name: str,
        globals_: dict[str, Any] | None = None,
        locals_: dict[str, Any] | None = None,
        fromlist: tuple[str, ...] = (),
        level: int = 0,
    ) -> Any:
        if name == "docx" or name.startswith("docx."):
            raise ModuleNotFoundError("docx")
        return real_import(name, globals_, locals_, fromlist, level)

    monkeypatch.setattr(builtins, "__import__", fake_import)


def test_render_docx_sin_python_docx_lanza_dependency_error(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """La ausencia del extra se traduce a un error accionable, no a un ``ModuleNotFoundError``."""
    _bloquear_python_docx(monkeypatch)

    with pytest.raises(ReportDependencyError, match=r"nikodym\[docx\]"):
        DocxReportRenderer.from_config(_config()).render(_bundle())


def test_docx_degrada_o_falla_sin_python_docx(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    """Sin el extra: ``fail_if_unavailable=False`` avisa y sigue; ``True`` re-lanza (como el PDF).

    Que falte un extra opcional no puede tumbar una corrida que ya calculó el modelo: el reporte
    principal se emite igual y el Word simplemente no sale, con aviso explícito.
    """
    _bloquear_python_docx(monkeypatch)
    bundle = _bundle()

    tolerante = DocxReportRenderer.from_config(_config(docx=DocxRenderConfig()))
    with pytest.warns(RuntimeWarning, match="python-docx no está disponible"):
        path = tolerante.write_docx_from_bundle(bundle, output_dir=str(tmp_path / "tolerante"))
    assert path is None
    assert not (tmp_path / "tolerante" / "scorecard_report.docx").exists()

    estricto = DocxReportRenderer.from_config(
        _config(docx=DocxRenderConfig(fail_if_unavailable=True))
    )
    with pytest.raises(ReportDependencyError, match="python-docx"):
        estricto.write_docx_from_bundle(bundle, output_dir=str(tmp_path / "estricto"))


def test_step_sin_python_docx_emite_el_reporte_igual(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    """Pedir ``docx`` sin el extra no rompe la corrida: sale el HTML y ``docx_path`` es None."""
    _bloquear_python_docx(monkeypatch)
    cfg = ReportConfig(output_dir=str(tmp_path), formats=("html", "docx"))
    study = _study_with_cards(cfg)

    with pytest.warns(RuntimeWarning, match="python-docx no está disponible"):
        result = ReportStep.from_config(cfg).execute(study, np.random.default_rng(1))

    assert result.docx_path is None
    assert (tmp_path / "scorecard_report.html").is_file()


# ─────────────────────── (c) import liviano (corre SIEMPRE) ───────────────────────


def test_import_report_no_arrastra_python_docx_por_subprocess() -> None:
    """``import nikodym.report`` NO importa python-docx (import perezoso, como WeasyPrint)."""
    code = (
        "import sys;"
        "sys.modules['docx'] = None;"
        "import nikodym.report as report;"
        "assert report.__name__ == 'nikodym.report';"
        "assert sys.modules.get('docx') is None"
    )
    subprocess.run([sys.executable, "-c", code], check=True)


def _study_with_cards(config: ReportConfig) -> Study:
    """``Study`` con las ocho cards requeridas y una tabla, listo para ``ReportStep``."""
    study = Study(NikodymConfig(report=config))
    study.run_context.lineage = _lineage()
    for domain, key in REPORT_REQUIRED_CARDS:
        study.artifacts.set(domain, key, {"summary": f"{domain}-card", "metric_sections": {}})
    study.artifacts.set(
        "model", "coefficients", pd.DataFrame({"feature": ["mora"], "beta": [1.25]})
    )
    return study
